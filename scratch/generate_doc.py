import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_document():
    doc = docx.Document()
    
    # Page Margins: Standard 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Text Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(55, 65, 81) # Charcoal gray

    # Document Title Block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("System Migration & Upgrade Report")
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(26)
    run_title.font.color.rgb = RGBColor(139, 92, 246) # Purple Accent
    p_title.paragraph_format.space_after = Pt(2)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("AI-Powered Semantic Search Engine upgrades: What, Why, and How")
    run_sub.italic = True
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(107, 114, 128)
    p_sub.paragraph_format.space_after = Pt(24)

    # Helper function to add structured headers
    def add_section_header(text, level=1):
        h = doc.add_heading(level=level)
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(17, 24, 39)
        if level == 1:
            run.font.size = Pt(18)
            h.paragraph_format.space_before = Pt(20)
            h.paragraph_format.space_after = Pt(8)
        elif level == 2:
            run.font.size = Pt(14)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
        return h

    # Helper function to add "What it was / What it is now" blocks
    def add_comparison_paragraph(what_was, what_is):
        p_was = doc.add_paragraph()
        r_was_lbl = p_was.add_run("What it was: ")
        r_was_lbl.bold = True
        r_was_lbl.font.color.rgb = RGBColor(220, 38, 38) # Soft Red
        p_was.add_run(what_was)
        
        p_is = doc.add_paragraph()
        r_is_lbl = p_is.add_run("What it is now: ")
        r_is_lbl.bold = True
        r_is_lbl.font.color.rgb = RGBColor(16, 185, 129) # Soft Green
        p_is.add_run(what_is)

    # 1. Introduction
    add_section_header("1. Document Overview")
    doc.add_paragraph(
        "This upgrade report provides a detailed breakdown of the recent changes made to the "
        "AI-Powered Semantic Search Engine. It covers transitions in the scraping layer, database storage "
        "architecture, user interface layout, local execution settings, and multi-format file support. "
        "For each change, the document explains what the feature was, what it has become, why the transition "
        "was necessary, and how it was implemented technically."
    )

    # 2. DuckDuckGo Web Search Migration
    add_section_header("2. DuckDuckGo Web Search Scraping Migration")
    
    add_comparison_paragraph(
        "The system sent standard HTTP GET requests to DuckDuckGo's static HTML URL (https://html.duckduckgo.com/html/?q=query).",
        "The system sends HTTP POST requests to the same URL, passing the search query as form data payload."
    )
    
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    doc.add_paragraph("Why it was changed:").bold = True
    doc.add_paragraph(
        "DuckDuckGo recently blocked or strictly rate-limited GET requests to their HTML search endpoint, "
        "returning an HTTP 400 Bad Request error. This caused the web search feature to fail. "
        "By switching to POST requests, we bypass this block because POST requests mimic browser form "
        "submissions, which DuckDuckGo allows."
    )
    
    doc.add_paragraph("How it was changed:").bold = True
    doc.add_paragraph(
        "In app.py, the search_ddg_html() function was modified. The requests.get() call was replaced "
        "with requests.post(), and the 'params' parameter was updated to 'data' to transmit the query "
        "parameters in the request body instead of the URL query string."
    )

    # 3. Local Database & Document Upload Transition
    add_section_header("3. Document Upload & Multi-Format Support")
    
    add_comparison_paragraph(
        "The local database was hardcoded in app.py as a static list of 5 articles. Uploading a document "
        "only supported PDF files, which wiped out any previously uploaded file (holding only one file at a time).",
        "The hardcoded static articles have been removed. The Local Database now supports uploading multiple files "
        "simultaneously, and parses PDF, DOCX, CSV, and TXT files, combining them into a single searchable index."
    )
    
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    doc.add_paragraph("Why it was changed:").bold = True
    doc.add_paragraph(
        "Users need to build their own custom knowledge base by uploading different types of documents. "
        "The previous setup restricted usage by only allowing one PDF file and mixing user data with hardcoded articles. "
        "Expanding support to PDF, DOCX, CSV, and TXT lets users index standard business and personal files."
    )
    
    doc.add_paragraph("How it was changed:").bold = True
    doc.add_paragraph(
        "1. Text Parsers: Added parse_docx() using the python-docx library to extract text from Word files (including tables) "
        "and parse_csv() using Python's built-in csv module to join spreadsheet cells.\n"
        "2. Multi-File Indexing: Replaced single-file overwrite logic in app.py. When a new file is uploaded, it is parsed and "
        "its chunks are assigned unique IDs based on a UUID ('filename_chunk_uuid'). If a file with the same name is uploaded, "
        "its old chunks are deleted first to avoid duplicates."
    )

    # 4. Storage Architecture (ChromaDB Server Integration)
    add_section_header("4. Database Storage Architecture")
    
    add_comparison_paragraph(
        "ChromaDB ran as an ephemeral (in-memory) database inside the Flask application process. "
        "Restarting the backend server immediately wiped out all uploaded documents.",
        "ChromaDB runs as a persistent service inside a separate Docker container. The backend "
        "connects to it via HTTP Client, and the database state is preserved across container restarts "
        "using a dedicated Docker volume."
    )
    
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    doc.add_paragraph("Why it was changed:").bold = True
    doc.add_paragraph(
        "A vector database must persist data so users do not have to re-upload files every time the "
        "server restarts. Separating the database into its own container matches containerized microservice "
        "best practices and keeps the Flask backend lightweight."
    )
    
    doc.add_paragraph("How it was changed:").bold = True
    doc.add_paragraph(
        "1. Docker Compose: Added a 'chroma' service running 'chromadb/chroma:latest' on port 8000 and mapped the "
        "named volume 'chroma-data' to '/chroma_data'.\n"
        "2. Environment Configuration: Configured environment variables (CHROMA_MODE, CHROMA_HOST, CHROMA_PORT) "
        "on the backend. On startup, app.py reads these variables. It connects via HttpClient if running in Docker, "
        "or falls back to writing locally to './chroma_db' using PersistentClient if running locally."
    )

    # 5. UI & UX Redesign (Two-Tab Console View)
    add_section_header("5. User Interface & Experience Redesign")
    
    add_comparison_paragraph(
        "The interface had three tabs: Web Search, Local Database (which searched the 5 hardcoded articles), "
        "and Document Search (which managed file uploads).",
        "The interface has only two tabs: Web Search and Local Database. The file manager and upload widget "
        "are integrated directly into the 'Local Database' tab, allowing users to configure and search their "
        "documents in one place."
    )
    
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    doc.add_paragraph("Why it was changed:").bold = True
    doc.add_paragraph(
        "The three-tab structure was redundant and confusing since the hardcoded database was deleted. "
        "Consolidating the workspace into a two-tab view matches standard search engine patterns: "
        "either search the Web or search your Local Database."
    )
    
    doc.add_paragraph("How it was changed:").bold = True
    doc.add_paragraph(
        "In App.tsx, the document search tab was removed. The file managers, file lists, and clear options "
        "were moved into the 'local' mode panel. If the user has no documents uploaded, "
        "the search input and submit buttons are disabled with a prompt asking them to upload files first."
    )

    # 6. Local Execution & Port Virtualization
    add_section_header("6. Local Port Virtualization")
    
    add_comparison_paragraph(
        "The Flask backend server started on port 5000, and the Vite client proxied requests to port 5000.",
        "The Flask backend server starts on port 5002, and the Vite client proxies requests to port 5002."
    )
    
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    doc.add_paragraph("Why it was changed:").bold = True
    doc.add_paragraph(
        "On macOS Monterey and newer, Apple runs Control Center / AirPlay Receiver on port 5000 by default. "
        "This causes port collision errors, preventing the backend from starting. Rebinding to port 5002 "
        "avoids this macOS system conflict."
    )
    
    doc.add_paragraph("How it was changed:").bold = True
    doc.add_paragraph(
        "1. Backend: Changed app.run() port to 5002 in app.py.\n"
        "2. Frontend: Changed fallback proxy target to http://127.0.0.1:5002 in vite.config.ts.\n"
        "3. Docker Compose: Remapped backend container port mapping from '5001:5000' to '5001:5002' "
        "and updated VITE_API_URL environment variable to point to port 5002."
    )

    doc.save("/Users/hamadalmeghewli/Code/semantic-search-demo/AI_Search_Engine_Documentation.docx")
    print("Updated migration report DOCX file created successfully.")

if __name__ == '__main__':
    create_document()
