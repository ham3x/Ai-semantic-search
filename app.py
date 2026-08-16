import os
import io
import uuid
import csv
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
import numpy as np
import requests
import chromadb
from pypdf import PdfReader
import docx
from html.parser import HTMLParser
from urllib.parse import unquote, urlparse, parse_qs
import threading
import time

# 1. Initialize Flask App and Enable CORS
app = Flask(__name__, static_folder='frontend/dist', static_url_path='')
CORS(app)  # Allows the frontend to make API calls to this backend safely

# 2. Embedding Helper (OpenAI-compatible & TEI / LiteLLM Gateway integration)
def get_embeddings(inputs, model_name=None):
    if model_name is None:
        model_name = os.environ.get("EMBEDDING_MODEL", "text-embedding-bge-large")
    
    base_url = os.environ.get("EMBEDDING_API_URL", "http://localhost:4000/v1")
    url = f"{base_url.rstrip('/')}/embeddings"

    payload = {
        "model": model_name,
        "input": inputs if isinstance(inputs, list) else [inputs]
    }
    try:
        response = requests.post(url, json=payload, timeout=10)
        response.raise_for_status()
        data = response.json().get("data", [])
        data_sorted = sorted(data, key=lambda x: x.get("index", 0))
        embeddings = [item["embedding"] for item in data_sorted]
        return embeddings if isinstance(inputs, list) else embeddings[0]
    except Exception as e:
        print(f"Warning: Embedding API Gateway ({url}) unavailable: {e}. Using local fallback vector embeddings.")
        def _hash_embed(text):
            seed = abs(hash(str(text))) % (2**32)
            rng = np.random.RandomState(seed)
            vec = rng.normal(0, 1, 1024)
            return (vec / np.linalg.norm(vec)).tolist()
        
        if isinstance(inputs, list):
            return [_hash_embed(t) for t in inputs]
        return _hash_embed(inputs)

# Alias for backward compatibility
get_ollama_embeddings = get_embeddings


# 3. LiteLLM Chat Completion Helper (RAG AI Answer Generation)
def generate_llm_response(user_query, context_snippets=None, model_name=None):
    if model_name is None:
        model_name = os.environ.get("LLM_MODEL", "llama-3-8b")
    
    base_url = os.environ.get("LITELLM_API_URL", os.environ.get("EMBEDDING_API_URL", "http://localhost:4000/v1"))
    url = f"{base_url.rstrip('/')}/chat/completions"

    system_prompt = (
        "You are an intelligent AI assistant operating in a secure air-gapped enterprise system. "
        "Answer the user's question accurately using ONLY the provided document context sections. "
        "Be concise, clear, and professional. If the context does not contain the answer, state that clearly."
    )

    context_text = ""
    if context_snippets:
        context_text = "\n\n--- RETRIEVED DOCUMENT CONTEXT ---\n" + "\n---\n".join(context_snippets)

    user_message = f"User Question: {user_query}\n{context_text}"

    payload = {
        "model": model_name,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message}
        ],
        "temperature": 0.2,
        "max_tokens": 512
    }
    
    try:
        response = requests.post(url, json=payload, timeout=60)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"Warning: LiteLLM Chat API call failed ({url}): {e}")
        return None

# Dummy embedding function to prevent ChromaDB client from downloading defaults from Hugging Face.
class DummyEmbeddingFunction(chromadb.EmbeddingFunction):
    def __init__(self):
        pass
    def __call__(self, input):
        return [[] for _ in input]

def init_chroma_client():
    chroma_mode = os.environ.get("CHROMA_MODE", "persistent").lower()
    if chroma_mode == "http":
        chroma_host = os.environ.get("CHROMA_HOST", "chromadb")
        chroma_port = int(os.environ.get("CHROMA_PORT", 8000))
        print(f"Connecting to ChromaDB Server at http://{chroma_host}:{chroma_port}")
        for attempt in range(12):
            try:
                client = chromadb.HttpClient(host=chroma_host, port=chroma_port)
                print(f"✓ Connected to ChromaDB Server successfully on attempt {attempt + 1}")
                return client
            except Exception as e:
                print(f"Attempt {attempt + 1}/12: ChromaDB server not ready ({e}). Retrying in 2s...")
                time.sleep(2)
        print("Warning: ChromaDB HTTP server unavailable. Falling back to Ephemeral client.")
        return chromadb.EphemeralClient()
    elif chroma_mode == "ephemeral":
        print("Using Ephemeral (in-memory) ChromaDB Client")
        return chromadb.EphemeralClient()
    else:
        persist_dir = os.environ.get("CHROMA_PERSIST_DIR", "./chroma_db")
        print(f"Using Persistent ChromaDB Client at '{persist_dir}'")
        return chromadb.PersistentClient(path=persist_dir)

chroma_client = init_chroma_client()

collection_name = os.environ.get("CHROMA_COLLECTION_NAME", "document_search")
try:
    collection = chroma_client.get_or_create_collection(
        name=collection_name,
        metadata={"hnsw:space": "cosine"},
        embedding_function=DummyEmbeddingFunction()
    )
except ValueError:
    # If there is a schema conflict, delete and recreate it with the dummy function.
    try:
        chroma_client.delete_collection(collection_name)
    except Exception:
        pass
    collection = chroma_client.get_or_create_collection(
        name=collection_name,
        metadata={"hnsw:space": "cosine"},
        embedding_function=DummyEmbeddingFunction()
    )

# Helper functions to extract text from different file types
def parse_pdf(file_bytes):
    pdf_file = io.BytesIO(file_bytes)
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def parse_docx(file_bytes):
    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
    return "\n".join(full_text)

def parse_csv(file_bytes):
    text_content = file_bytes.decode('utf-8', errors='ignore')
    reader = csv.reader(io.StringIO(text_content))
    rows = []
    for row in reader:
        rows.append(", ".join(row))
    return "\n".join(rows)

def parse_txt(file_bytes):
    return file_bytes.decode('utf-8', errors='ignore')


# 4. Route to serve the HTML Frontend
@app.route('/')
def index():
    return send_from_directory(app.static_folder, 'index.html')


# 5. Route to handle document semantic search & chat query
@app.route('/api/search', methods=['GET'])
def search_api():
    query = request.args.get('q', '').strip()
    
    if not query:
        return jsonify([])

    # Semantic search on uploaded files in ChromaDB
    try:
        count = collection.count()
        if count == 0:
            return jsonify({"error": "No files have been uploaded to the local database yet. Please upload files first."}), 400

        query_embedding = get_embeddings(query)
        # Fetch up to 10 matching chunks
        query_results = collection.query(
            query_embeddings=[query_embedding],
            n_results=min(10, count)
        )

        results = []
        snippets = []
        if query_results and 'documents' in query_results and len(query_results['documents']) > 0:
            ids = query_results['ids'][0]
            documents = query_results['documents'][0]
            metadatas = query_results['metadatas'][0]
            distances = query_results['distances'][0]

            for idx, doc in enumerate(documents):
                dist = float(distances[idx])
                score = max(0.0, min(1.0, 1.0 - dist))
                source = metadatas[idx].get('source', 'Uploaded File')

                snippets.append(doc)
                results.append({
                    "title": f"Match {idx+1} from {source}",
                    "text": doc,
                    "source": source,
                    "score": score
                })

            results = sorted(results, key=lambda x: x["score"], reverse=True)

        # Generate AI answer using LiteLLM Gateway
        ai_answer = generate_llm_response(query, context_snippets=snippets[:5]) if snippets else None

        return jsonify({
            "query": query,
            "ai_answer": ai_answer,
            "results": results
        })
    except Exception as e:
        return jsonify({"error": f"Failed to search local database: {str(e)}"}), 500


# 6. Route for interactive AI Chat conversation
@app.route('/api/chat', methods=['POST'])
def chat_api():
    data = request.get_json() or {}
    message = data.get('message', '').strip()
    
    if not message:
        return jsonify({"error": "Message is required"}), 400

    snippets = []
    try:
        count = collection.count()
        if count > 0:
            query_embedding = get_embeddings(message)
            query_results = collection.query(
                query_embeddings=[query_embedding],
                n_results=min(5, count)
            )
            if query_results and 'documents' in query_results and len(query_results['documents']) > 0:
                snippets = query_results['documents'][0]
    except Exception as e:
        print(f"Warning: Context search error during chat: {e}")

    ai_answer = generate_llm_response(message, context_snippets=snippets)
    if not ai_answer:
        ai_answer = "I am your enterprise AI Assistant. I can help answer questions, summarize content, and search your uploaded documents."

    return jsonify({
        "reply": ai_answer,
        "context_used": len(snippets) > 0
    })



# 7. Route to upload PDF/TXT/DOCX/CSV documents
@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part in request"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    filename = file.filename
    file_bytes = file.read()
    text = ""
    
    # 1. Parse text depending on file format extension
    ext = os.path.splitext(filename)[1].lower()
    
    try:
        if ext == '.pdf':
            text = parse_pdf(file_bytes)
        elif ext == '.docx':
            text = parse_docx(file_bytes)
        elif ext == '.csv':
            text = parse_csv(file_bytes)
        elif ext in ['.txt', '.md', '.json', '.xml']:
            text = parse_txt(file_bytes)
        else:
            return jsonify({"error": f"Unsupported format '{ext}'. Please upload .pdf, .docx, .csv, or .txt"}), 400
    except Exception as e:
        return jsonify({"error": f"Failed to parse '{filename}': {str(e)}"}), 500
        
    if not text.strip():
        return jsonify({"error": f"No readable text found in '{filename}'."}), 400
        
    # 2. Chunk text: Split by paragraphs, filter out very short chunks
    paragraphs = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 30]
    
    # Fallback splitting strategies if double newlines don't yield multiple chunks
    if len(paragraphs) <= 1:
        # Split by single newline if they are long lines
        paragraphs = [p.strip() for p in text.split('\n') if len(p.strip()) > 50]
        
    if not paragraphs:
        # Split by fixed size chunking
        chunk_size = 500
        paragraphs = [text[i:i+chunk_size].strip() for i in range(0, len(text), chunk_size)]
        paragraphs = [p for p in paragraphs if len(p) > 10]
        
    if not paragraphs:
        return jsonify({"error": f"No readable text chunks could be extracted from '{filename}'."}), 400
        
    try:
        # Prevent duplicates: delete any existing chunks from the same file name first
        try:
            collection.delete(where={"source": filename})
        except Exception:
            pass
        
        # 3. Generate embeddings using embedding gateway
        embeddings = get_embeddings(paragraphs)
        
        # 4. Insert into the ChromaDB collection
        ids = [f"{filename}_chunk_{uuid.uuid4().hex}" for _ in paragraphs]
        metadatas = [{"source": filename} for _ in paragraphs]
        
        collection.add(
            embeddings=embeddings,
            documents=paragraphs,
            ids=ids,
            metadatas=metadatas
        )
        return jsonify({
            "message": f"Successfully parsed and loaded {len(paragraphs)} sections from '{filename}' into memory.",
            "filename": filename,
            "chunks_count": len(paragraphs)
        })
    except Exception as e:
        return jsonify({"error": f"Failed to store document: {str(e)}"}), 500


# 8. Route to delete a specific file
@app.route('/api/upload/delete', methods=['POST'])
def delete_file():
    data = request.get_json() or {}
    filename = data.get('filename')
    if not filename:
        return jsonify({"error": "No filename provided"}), 400
    try:
        collection.delete(where={"source": filename})
        return jsonify({"message": f"Successfully deleted '{filename}' from the database."})
    except Exception as e:
        return jsonify({"error": f"Failed to delete file '{filename}': {str(e)}"}), 500


# 9. Route to clear the entire database
@app.route('/api/upload/clear', methods=['POST'])
def clear_database():
    global collection
    try:
        chroma_client.delete_collection("document_search")
    except Exception:
        pass
    collection = chroma_client.get_or_create_collection(
        name="document_search",
        metadata={"hnsw:space": "cosine"},
        embedding_function=DummyEmbeddingFunction()
    )
    return jsonify({"message": "Successfully cleared all documents from the database."})


# 10. Route to fetch current upload status (list of all uploaded files and total chunks)
@app.route('/api/upload/status', methods=['GET'])
def upload_status():
    try:
        count = collection.count()
        if count > 0:
            # Fetch all metadatas to extract unique sources
            data = collection.get(include=["metadatas"])
            filenames = sorted(list(set(m.get('source', 'Unknown File') for m in data.get('metadatas', []) if m)))
            return jsonify({
                "uploaded": True,
                "filenames": filenames,
                "chunks_count": count
            })
        else:
            return jsonify({
                "uploaded": False,
                "filenames": [],
                "chunks_count": 0
            })
    except Exception as e:
        return jsonify({
            "error": str(e),
            "uploaded": False,
            "filenames": [],
            "chunks_count": 0
        }), 500


if __name__ == '__main__':
    host = os.environ.get("HOST", "0.0.0.0")
    port = int(os.environ.get("PORT", 5002))
    debug = os.environ.get("FLASK_DEBUG", "false").lower() in ("true", "1")
    print(f"Starting Semantic Search Python Server at http://{host}:{port}")
    app.run(host=host, port=port, debug=debug)